import os
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain.schema import HumanMessage
from pdf2image import convert_from_path
import pytesseract
import json


# Load environment variables
load_dotenv()
groq_api_key = os.getenv('GROQ_API_KEY')

from docx import Document

def organize_unmapped_technologies(parsed_response):
    # Get all technologies from categories
    all_technologies = set()
    for category in parsed_response['technology_categories']:
        all_technologies.update(category['technologies'])
    
    # Get mapped technologies from projects
    mapped_technologies = set()
    for company_projects in parsed_response['project_technology_links']:
        for project in company_projects['projects']:
            mapped_technologies.update(project['technologies'])
    
    # Find unmapped technologies
    unmapped_technologies = all_technologies - mapped_technologies
    
    if unmapped_technologies:
        # Add Additional Experience company and Projects project
        additional_experience = {
            "company": "Additional Experience",
            "projects": [{
                "name": "Additional Projects",
                "technologies": list(unmapped_technologies),
                "description": "Additional technical experience and tools"
            }]
        }
        
        # Add to companies list if not exists
        if not any(company['name'] == "Additional Experience" 
                  for company in parsed_response['companies']):
            parsed_response['companies'].append({
                "name": "Additional Experience",
                "startDate": "",
                "endDate": "Present"
            })
        
        # Add to project_technology_links
        parsed_response['project_technology_links'].append(additional_experience)
    
    return parsed_response
    
def link_projects_to_technologies(chat, text_content, technologies, companies):
    project_tech_prompt = [
        HumanMessage(content=f"""
        For each company and their projects in the resume, identify which technologies from the verified list were used.
        
        Resume text:
        {text_content}
        
        Verified technologies:
        {technologies}
        
        Companies:
        {companies}
        
        Return as JSON array:
        [{{
            "company": "company name",
            "projects": [{{
                "name": "project name",
                "technologies": ["tech1", "tech2"],
                "description": "project description"
            }}]
        }}]
        """)
    ]
    
    response = chat.invoke(project_tech_prompt)
    return json.loads(response.content)


def verify_technologies(chat, text_content, technologies):
    verification_prompt__ = [
        HumanMessage(content=f"""
        Compare these technologies with the resume text and verify their presence.
        Remove any technology that cannot be found in the original text.
        
        Resume text:
        {text_content}
        
        Technologies to verify:
        {technologies}
        
        Return only a JSON array of verified technologies that are actually mentioned in the resume.
        Format: ["Technology1", "Technology2", ...]
        """)
    ]
    verification_prompt = [
        HumanMessage(content=f"""
        STRICTLY verify these technologies against the resume text.
        A technology should ONLY be included if it is EXPLICITLY mentioned in the resume text.
        Do not include:
        - Similar technologies
        - Related tools
        - Inferred technologies
        - Technologies that might be used but aren't explicitly mentioned
        
        Resume text:
        {text_content}
        
        Technologies to verify:
        {technologies}
        
        Rules:
        1. Only return technologies that appear WORD FOR WORD in the resume
        2. Maintain exact spelling and capitalization as found in the resume
        3. Do not add any technologies that aren't explicitly listed
        4. Do not make assumptions about related technologies
        
        Return as JSON array: ["Technology1", "Technology2", ...]
        """)
     ]
    response = chat.invoke(verification_prompt)
    return json.loads(response.content)

def categorize_verified_technologies(chat, verified_technologies):
    categorization_prompt = [
        HumanMessage(content=f"""
        Categorize these verified technologies according to these rules:
        - Hardware items → "Hardware & IoT"
        - Programming languages → "Programming Languages"
        - Database technologies → "Data Storage"
        - Cloud services → "Cloud & Infrastructure"
        - Development tools → "Development Tools"
        - AI/ML technologies → "Machine Learning & AI"
        - Web frameworks → "Web Technologies"
        - NLP tools → "NLP"
        - Operating systems → "Operating Systems"
        - DevOps tools → "DevOps"

        Technologies to categorize:
        {verified_technologies}

        Return as JSON with categories as keys and technology arrays as values:
        {{
            "category1": ["tech1", "tech2"],
            "category2": ["tech3", "tech4"]
        }}
        """)
    ]
    
    response = chat.invoke(categorization_prompt)
    return json.loads(response.content)

def normalize_technologies(chat, categorized_technologies):
    normalization_prompt = [
        HumanMessage(content=f"""
        Clean and normalize this technology list by:
        1. Combining duplicates (e.g., "MS SQL"/"MSSQL" → "Microsoft SQL Server")
        2. Standardizing names (e.g., "React.js"/"ReactJS" → "React")
        3. Removing non-technical items (languages, etc.)
        4. Ensuring each technology appears once

        Categorized technologies:
        {json.dumps(categorized_technologies, indent=2)}

        Return the cleaned data in the same JSON structure.
        """)
    ]
    
    response = chat.invoke(normalization_prompt)
    return json.loads(response.content)
    


def categorize_technologies(technologies_list):
    # Second LLM call specifically for categorization
    chat = ChatGroq(
        api_key=groq_api_key,
        model_name="mixtral-8x7b-32768"
    )
    
    categorization_prompt = [
        HumanMessage(content=f"""
        Categorize these technologies into appropriate groups and return as JSON:
        Input technologies: {', '.join(technologies_list)}
        
        Format the response as an array of categories:
        [
            {
                "category_name": "Programming Languages",
                "technologies": ["C++", "C#", "Java", etc]
            },
            {
                "category_name": "Web Technologies",
                "technologies": ["React", "Angular", "ASP.NET", etc]
            },
            {
                "category_name": "Databases",
                "technologies": ["SQL Server", "MongoDB", etc]
            },
            {
                "category_name": "Cloud & Infrastructure",
                "technologies": ["AWS", "Azure", etc]
            }
        ]
        """)
    ]
    
    response = chat.invoke(categorization_prompt)
    return response.content

def extract_word_text(doc_path):
    doc = Document(doc_path)
    text_content = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text_content.append(paragraph.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text_content.append(' '.join(row_text))
           
    print(f"Text: {text_content}")
    
    return '\n'.join(text_content)
    
def extract_pdf_text(pdf_path):
    try:
        # Convert PDF to images
        images = convert_from_path(pdf_path)
        
        # Extract text from all pages
        text_content = []
        for image in images:
            text = pytesseract.image_to_string(image)
            text_content.append(text)
        
        print(f"Text: {text_content}")
        return " ".join(text_content)
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""

def process_resume(pdf_path):
    # Initialize Groq client
    
    chat = ChatGroq(
		api_key=groq_api_key,
		model_name="mixtral-8x7b-32768"  
	)
    #model_name="mixtral-8x7b-32768"  # Use this supported model instead
    #deepseek-r1-distill-llama-70b
    
    # Extract text from PDF
    #text_content = extract_pdf_text(pdf_path)
    text_content = extract_word_text(pdf_path)
    
    # Create prompt
    messages = [
    HumanMessage(content=f"""
    Extract the following information from this resume and format as JSON:
    - companies: list of company experiences with structure:
    [{{
        "name": "company name",
        "startDate": "MM/YYYY",
        "endDate": "MM/YYYY or Present"
    }}]
    - projects: list of projects with client names
    - education: educational background
    - technologies: list of technologies used
    - titles: job titles held

    For dates, use the format "MM/YYYY". If it's a current position, use "Present" for endDate.
    Ensure all companies have both start and end dates.
    
    Resume content:
    {text_content}
    """)
    ]

    
    # Get response
    response = chat.invoke(messages)
    parsed_response = json.loads(response.content)
    
    if 'technologies' in parsed_response:
        verified_techdata = verify_technologies(chat, text_content, parsed_response['technologies'])
        print("Verify Tech Data:", verified_techdata)
        tech_messages_ = [
            HumanMessage(content=f"""
            Strictly categorize ONLY the following technologies into appropriate groups and DO NOT add any technologies that are not in this list then only return as JSON array:
            Technologies to categorize: {verified_techdata}
            
            Return only a valid JSON array without any comments like this:
            [
                {{"categoryName": "Programming Languages",
                  "technologies": ["Tech1", "Tech2"]}},
                {{"categoryName": "Web Technologies",
                  "technologies": ["Tech3", "Tech4"]}}
            ]
            """)
        ]

        #categorize_technologies = categorize_verified_technologies(chat, verified_techdata)
        categorize_technologies = chat.invoke(tech_messages_)
        final_technologies = categorize_technologies.content.strip()
        final_technologies = final_technologies.replace('\n', '')
        print("Categorize:", final_technologies)
        # Remove any leading/trailing whitespace and newlines
        #json_content = json_content.replace('\n', '')
        #final_technologies = normalize_technologies(chat, json_content)
        print("Finalized:", final_technologies)
        
        tech_messages = [
            HumanMessage(content=f"""
            Categorize these technologies into appropriate groups and return as JSON array:
            Technologies: {', '.join(parsed_response['technologies'])}
            
            Return only a valid JSON array without any comments like this:
            [
                {{"categoryName": "Programming Languages",
                  "technologies": ["Tech1", "Tech2"]}},
                {{"categoryName": "Web Technologies",
                  "technologies": ["Tech3", "Tech4"]}}
            ]
            """)
        ]
        
        project_links = link_projects_to_technologies(
            chat, 
            text_content, 
            verified_techdata, 
            parsed_response['companies']
        )
        print("Project Technologies:", project_links)
        
        
        try:
            #tech_response = chat.invoke(tech_messages)
            #cleaned_response = tech_response.content.strip()
            # Print for debugging
            #print("Raw response:", cleaned_response)
            #parsed_response['technology_categories'] = json.loads(cleaned_response)
            parsed_response['technology_categories'] = json.loads(final_technologies)
            parsed_response['project_technology_links'] = project_links
            parsed_response = organize_unmapped_technologies(parsed_response)
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {e}")
            parsed_response['technology_categories'] = []
            parsed_response['project_technology_links'] = []
        
        except Exception as e:
            print(f"Error categorizing technologies: {e}")
            parsed_response['technology_categories'] = []
            parsed_response['project_technology_links'] = []
        
    
    return json.dumps(parsed_response, indent=2)
    
    #return response.content
